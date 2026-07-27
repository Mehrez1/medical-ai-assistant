"""
Interface Streamlit - Intégration RAG + Memory + Error Handling
"""

import streamlit as st
import json
import os
import sys
from datetime import datetime
from dotenv import load_dotenv
from pathlib import Path

# Charger les variables d'environnement depuis .env au chemin correct
project_root = Path(__file__).parent.parent
load_dotenv(project_root / ".env")

# 🔑 CHARGER CLÉ API UNE FOIS ET CONSERVER TOUJOURS
if "api_key_loaded" not in st.session_state:
    groq_api_key = os.getenv("GROQ_API_KEY")
    st.session_state.groq_api_key = groq_api_key
    st.session_state.api_key_loaded = True
    st.session_state.api_key_persistent = True  # Marquer comme persistant
    
    # Log
    if groq_api_key:
        print("✅ Clé API Groq chargée et persistante pour la durée de la session")
    else:
        print("⚠️ ATTENTION: Clé API Groq non trouvée!")

# Ajouter le répertoire parent au path pour les imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Import des modules Phase 1
from src.memory import ConversationMemory
from src.pipelines import RAGPipeline
from src.error_handler import get_error_handler, NoSymptomsDetected
from src.explainability import get_explainability
from src.tools import MedicalTools
from src.agent import MedicalOrientationAgent
from src.rag import MedicalRAG

# Configuration
st.set_page_config(
    page_title="🏥 Assistant Médical",
    page_icon="🏥",
    layout="wide"
)

# CSS
st.markdown("""
<style>
    .urgence-banal { background-color: #d4edda; color: #155724; padding: 15px; border-radius: 5px; font-weight: bold; }
    .urgence-normal { background-color: #fff3cd; color: #856404; padding: 15px; border-radius: 5px; font-weight: bold; }
    .urgence-urgent { background-color: #f8d7da; color: #721c24; padding: 15px; border-radius: 5px; font-weight: bold; }
    .urgence-extreme { background-color: #f5c6cb; color: #721c24; padding: 15px; border-radius: 5px; font-weight: bold; border: 2px solid red; }
</style>
""", unsafe_allow_html=True)

# Charger la base de données médicale
def load_medical_data():
    data_path = "data/medical_data.json"
    if os.path.exists(data_path):
        with open(data_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

# Initialiser les modules Phase 1
@st.cache_resource
def initialize_modules():
    """Initialise une fois les modules (cached)"""
    error_handler = get_error_handler()
    memory = ConversationMemory(max_history=20)
    rag_pipeline = RAGPipeline()
    agent = MedicalOrientationAgent()
    explainability = get_explainability()
    return {
        "error_handler": error_handler,
        "memory": memory,
        "rag_pipeline": rag_pipeline,
        "agent": agent,
        "explainability": explainability
    }

medical_data = load_medical_data()
modules = initialize_modules()
error_handler = modules["error_handler"]
memory = modules["memory"]
rag_pipeline = modules["rag_pipeline"]
agent = modules["agent"]
explainability = modules["explainability"]

# TITRE
st.title("🏥 Assistant d'Orientation Médicale")
st.markdown("""
Décrivez vos symptômes et l'assistant vous orientera vers le bon professionnel.
⚠️ **Ce n'est PAS un diagnostic médical** - Consultez toujours un professionnel.
""")

# Initialiser session
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = []  # Historique du chat conversationnel
if "results_shown" not in st.session_state:
    st.session_state.results_shown = False
if "results" not in st.session_state:
    st.session_state.results = None
if "selected_provider" not in st.session_state:
    st.session_state.selected_provider = "groq"  # Groq par défaut
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = []  # Historique du chat conversationnel

# SIDEBAR
with st.sidebar:
    st.header("🎯 Contrôles")
    if st.button("🔄 Réinitialiser", use_container_width=True):
        st.session_state.chat_messages = []
        st.session_state.results_shown = False
        st.rerun()
    
    # Sélecteur LLM
    st.divider()
    st.subheader("⚙️ Configuration LLM")
    from src.llm_factory import LLMFactory, set_llm_provider
    
    llm_models = LLMFactory.get_available_models()
    model_keys = list(llm_models.keys())
    
    # Utiliser le provider sauvegardé dans session_state, sinon Groq
    current_provider = st.session_state.selected_provider
    default_index = model_keys.index(current_provider) if current_provider in model_keys else 0
    
    selected_provider = st.selectbox(
        "Choisir le modèle LLM:",
        options=model_keys,
        index=default_index,
        format_func=lambda x: llm_models[x]["name"],
        help="Vous pouvez changer le modèle LLM selon vos préférences",
        key="llm_selector"
    )
    
    # Mettre à jour session_state si l'utilisateur a changé de provider
    if selected_provider != st.session_state.selected_provider:
        st.session_state.selected_provider = selected_provider
    
    # APPLIQUER LE CHANGEMENT DE LLM A L'AGENT (avec gestion d'erreur)
    try:
        agent.set_provider(selected_provider)
    except ValueError as e:
        error_msg = str(e)
        st.warning(f"⚠️ {error_msg}")
        # Revenir au provider par défaut
        st.session_state.selected_provider = "groq"
        agent.set_provider("groq")
        selected_provider = "groq"
    
    # Afficher infos du modèle sélectionné
    model_info = llm_models[selected_provider]
    with st.container():
        st.write(f"**Vitesse:** {model_info['speed']}")
        st.write(f"**Qualité:** {model_info['quality']}")
        st.write(f"**Coût:** {model_info['cost']}")
        st.caption(model_info['description'])

# INTRO
st.markdown("""
**🚀 Un agent IA unique qui:**
- ✅ Détecte automatiquement les symptômes et maladies
- ✅ Répond à n'importe quelle question (médicale ou générale)
- ✅ Analyse vos fichiers (texte, PDF, rapports médicaux, etc.)

Posez une question directement ou uploadez un fichier - l'agent s'adapte automatiquement! 🤖
""")
st.divider()

# ═══════════════════════════════════════════════════════════════════
# SECTION: CHAT LIBRE CONVERSATIONNEL ✨ (NOUVEAU!)
# ═══════════════════════════════════════════════════════════════════

st.divider()
st.subheader("💬 Chat Libre - Posez une Question à l'IA")
st.caption("L'agent IA peut répondre à n'importe quelle phrase (médicale ou générale)")

# Afficher l'historique du chat
for message in st.session_state.chat_messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Champ de saisie du chat
user_input = st.chat_input(
    placeholder="Écrivez votre message... (ex: Bonjour, comment allez-vous? ou J'ai une douleur à la poitrine)",
    key="chat_input"
)

if user_input:
    # Ajouter le message de l'utilisateur à l'historique
    st.session_state.chat_messages.append({
        "role": "user",
        "content": user_input
    })
    
    # Afficher le message utilisateur
    with st.chat_message("user"):
        st.markdown(user_input)
    
    # Générer réponse de l'agent IA UNIQUE
    with st.chat_message("assistant"):
        with st.spinner("🤔 L'IA traite votre demande..."):
            try:
                # Appeler l'agent unique qui s'adapte automatiquement
                result = agent.process_user_input(user_input)
                
                # Afficher la réponse
                st.markdown(result["response"])
                
                # Ajouter la réponse à l'historique
                st.session_state.chat_messages.append({
                    "role": "assistant",
                    "content": result["response"]
                })
                
            except Exception as e:
                error_msg = f"❌ Erreur: {str(e)}\n\nVeuillez réessayer."
                st.error(error_msg)
                
                # Ajouter le message d'erreur à l'historique
                st.session_state.chat_messages.append({
                    "role": "assistant",
                    "content": error_msg
                })

# ═══════════════════════════════════════════════════════════════════
# SECTION: ANALYSE DE FICHIERS 📁 (NOUVEAU!)
# ═══════════════════════════════════════════════════════════════════

st.divider()
st.subheader("📁 Analyser un Fichier")
st.caption("Uploadez n'importe quel fichier (texte, PDF, Word, Excel, images, etc.) pour l'analyser automatiquement")

# Créer deux colonnes pour upload et question
col_upload, col_question = st.columns([1, 2])

with col_upload:
    uploaded_file = st.file_uploader(
        "Choisir un fichier:",
        type=["txt", "md", "csv", "pdf", "jpg", "jpeg", "png", "docx", "xlsx", "xls", "doc", "rtf", "odt"],
        help="Formats supportés: Texte, PDF, Word, Excel, Images, etc."
    )

with col_question:
    file_question = st.text_input(
        "Question sur le fichier (optionnel):",
        placeholder="Ex: Analyser ce rapport médical ou Qu'est-ce qui est anormal?",
        key="file_question"
    )

if uploaded_file is not None:
    # Initialiser session pour fichier
    if "file_analyzed" not in st.session_state:
        st.session_state.file_analyzed = False
    
    # Lire le contenu du fichier
    try:
        file_bytes = uploaded_file.read()
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        # TEXTE SIMPLE (TXT, MD, CSV, RTF, etc.)
        if uploaded_file.type in ["text/plain", "text/markdown", "text/csv", "text/rtf"] or file_extension in ["txt", "md", "csv", "rtf"]:
            # Essayer plusieurs encodages en ordre de priorité
            encodings_to_try = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252', 'utf-16']
            file_content = None
            
            for encoding in encodings_to_try:
                try:
                    file_content = file_bytes.decode(encoding)
                    break
                except (UnicodeDecodeError, LookupError):
                    continue
            
            # Si aucun encodage ne fonctionne, utiliser utf-8 avec les erreurs ignorées
            if file_content is None:
                file_content = file_bytes.decode('utf-8', errors='ignore')
            
            file_type_label = "📝 Texte"
        
        # PDF
        elif uploaded_file.type == "application/pdf" or file_extension == "pdf":
            try:
                import PyPDF2
                from io import BytesIO
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
                file_content = "\n".join([page.extract_text() for page in pdf_reader.pages])
                file_type_label = "📄 PDF"
            except ImportError:
                file_content = f"[PDF détecté mais PyPDF2 non installé. Veuillez installer: pip install PyPDF2]"
                file_type_label = "📄 PDF (non supporté)"
            except Exception as e:
                file_content = f"[Erreur PDF: {str(e)}]"
                file_type_label = "📄 PDF (erreur)"
        
        # WORD DOCUMENT (DOCX et DOC)
        elif "wordprocessingml" in uploaded_file.type or file_extension in ["docx", "doc"]:
            file_content = None
            file_type_label = "📘 Word"
            
            # Essayer avec python-docx (pour .docx et .doc modernes)
            try:
                from docx import Document
                from io import BytesIO
                doc = Document(BytesIO(file_bytes))
                file_content = "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                file_content = f"[Fichier Word détecté. Installation de python-docx nécessaire.]\npip install python-docx"
            except Exception as e_docx:
                # Fallback: essayer de convertir .doc en texte avec doc2docx
                try:
                    from doc2docx import convert
                    import tempfile
                    from pathlib import Path
                    
                    with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                        tmp.write(file_bytes)
                        tmp_path = tmp.name
                    
                    output_path = tmp_path.replace('.doc', '_converted.docx')
                    convert(tmp_path, output_path)
                    
                    from docx import Document
                    doc = Document(output_path)
                    file_content = "\n".join([para.text for para in doc.paragraphs])
                    
                    # Cleanup
                    import os
                    os.remove(tmp_path)
                    os.remove(output_path)
                    
                except ImportError:
                    file_content = f"[Ancien format .doc détecté. Installation de doc2docx nécessaire.]\npip install doc2docx"
                except Exception as e_doc2:
                    file_content = f"[Impossible de lire le fichier Word: {str(e_docx)}]"
        
        # EXCEL (XLSX, XLS)
        elif "spreadsheetml" in uploaded_file.type or file_extension in ["xlsx", "xls"]:
            try:
                import pandas as pd
                excel_file = pd.ExcelFile(uploaded_file)
                sheets_content = []
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
                    sheets_content.append(f"\n=== Feuille: {sheet_name} ===\n{df.to_string()}")
                file_content = "\n".join(sheets_content)
                file_type_label = "📊 Excel"
            except ImportError:
                file_content = f"[Fichier Excel détecté. Installation de openpyxl et pandas nécessaire pour l'extraction.]\n\nNom: {uploaded_file.name}\nTaille: {uploaded_file.size} bytes"
                file_type_label = "📊 Excel"
            except Exception as e:
                file_content = f"[Fichier Excel mais erreur: {str(e)}]\n\nNom: {uploaded_file.name}\nTaille: {uploaded_file.size} bytes"
                file_type_label = "📊 Excel"
        
        # IMAGES
        elif uploaded_file.type in ["image/jpeg", "image/png", "image/jpg"] or file_extension in ["jpg", "jpeg", "png"]:
            file_content = f"[Image détectée]\nNom: {uploaded_file.name}\nTaille: {uploaded_file.size} bytes\nNote: Décrivez ce que vous voyez dans l'image pour une analyse complète."
            file_type_label = "🖼️ Image"
        
        # AUTRES FORMATS
        else:
            # Essayer de le lire comme texte générique
            try:
                file_content = file_bytes.decode('utf-8', errors='ignore')
                file_type_label = f"📄 {file_extension.upper()}"
            except:
                file_content = f"[Fichier: {uploaded_file.name}]\nTaille: {uploaded_file.size} bytes\nFormat: {file_extension}\n\nLe fichier a été détecté mais ne peut pas être entièrement décodé. Veuillez vérifier le format."
                file_type_label = f"📦 {file_extension.upper()}"
        
        # Afficher info du fichier
        st.info(f"{file_type_label} | {uploaded_file.name} | {uploaded_file.size} bytes")
        
        # Bouton d'analyse
        if st.button("🔍 Analyser le Fichier", use_container_width=True, type="primary"):
            with st.spinner("🤔 L'agent analyse votre fichier..."):
                try:
                    # Utiliser l'agent UNIQUE qui s'adapte automatiquement
                    result = agent.process_user_input(
                        user_input=file_content,
                        is_file_content=True,
                        file_name=uploaded_file.name
                    )
                    
                    st.success("✅ Analyse complétée!")
                    st.divider()
                    
                    # AFFICHER LES RÉSULTATS EN FONCTION DU TYPE DÉTECTÉ
                    if result["is_medical"]:
                        # ═══════════════════════════════════════════════════════════════
                        # MODE: ANALYSE MÉDICALE COMPLÈTE
                        # ═══════════════════════════════════════════════════════════════
                        st.subheader("📋 Analyse Médicale Complète")
                        
                        # Afficher analyse LLM
                        if result.get("response"):
                            provider_display = st.session_state.get("selected_provider", "groq").upper()
                            st.info(f"**Analyse par {provider_display}:**\n\n{result.get('response')}")
                            st.divider()
                        
                        # Afficher symptômes détectés
                        col_symp, col_context = st.columns([2, 1])
                        with col_symp:
                            st.write("**Symptômes Détectés:**")
                            if result.get("symptoms"):
                                for s in result.get("symptoms", []):
                                    st.write(f"• {s}")
                            else:
                                st.write("Aucun symptôme clair détecté")
                        
                        with col_context:
                            st.write("**Contexte:**")
                            st.write(f"- Sévérité: {result.get('severity', 'N/A').upper()}")
                        
                        # Afficher urgence
                        if result.get("urgency"):
                            urgency_data = result.get("urgency", {})
                            urgency_level = urgency_data.get("level", "NORMAL")
                            
                            urgency_colors = {
                                "BANAL": "urgence-banal",
                                "NORMAL": "urgence-normal",
                                "URGENT": "urgence-urgent",
                                "EXTRÊME URGENCE": "urgence-extreme"
                            }
                            
                            urgency_class = urgency_colors.get(urgency_level, "urgence-normal")
                            
                            col_res1, col_res2 = st.columns(2)
                            
                            with col_res1:
                                st.markdown(f"""<div class="{urgency_class}">
                                🚨 {urgency_level}
                                </div>""", unsafe_allow_html=True)
                                st.write(f"**Recommandation:**\n{urgency_data.get('recommendation', 'Consulter un médecin')}")
                            
                            with col_res2:
                                specialist_data = result.get("specialist", {})
                                st.info(f"**Spécialiste Recommandé:**\n{specialist_data.get('specialist', 'Médecin généraliste')}")
                    else:
                        # ═══════════════════════════════════════════════════════════════
                        # MODE: RÉPONSE LIBRE (pas de cas médical détecté)
                        # ═══════════════════════════════════════════════════════════════
                        st.subheader("💬 Réponse de l'Agent")
                        st.markdown(result.get("response", ""))
                    
                    # Ajouter au chat
                    st.session_state.chat_messages.append({
                        "role": "user",
                        "content": f"📁 Fichier: {uploaded_file.name}"
                    })
                    st.session_state.chat_messages.append({
                        "role": "assistant",
                        "content": result.get("response", "")
                    })
                    
                    st.session_state.file_analyzed = True
                    
                except Exception as e:
                    st.error(f"❌ Erreur lors de l'analyse: {str(e)}")
                    print(f"Erreur détaillée: {str(e)}")  # Pour debug
    
    except Exception as e:
        st.error(f"❌ Erreur lors de la lecture du fichier: {str(e)}")
